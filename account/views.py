from django.shortcuts import render, redirect
from .forms import SignUpForm, LoginForm, ClaseForm
from django.contrib.auth import authenticate, login
from django.contrib import messages
from .models import Unidad, Contenido, Clase
from django.http import JsonResponse
from .models import Clase, User
from docx import Document
from account.models import Semestre, Materia, Unidad, Contenido
from docx import Document
import docx

import re
from .forms import DocumentForm
import tabula
import camelot
from django.http import HttpResponseRedirect
from django.urls import reverse
import os
from django.conf import settings

import pdfplumber
from io import BytesIO
from unidecode import unidecode


# Create your views here.


def index(request):
    return render(request, 'index.html')


def register(request):
    msg = None
    if request.method == 'POST':
        form = SignUpForm(request.POST)
        if form.is_valid():
            user = form.save()
            msg = 'user created'
            return redirect('login_view')
        else:
            msg = 'form is not valid'
    else:
        form = SignUpForm()
    return render(request,'register.html', {'form': form, 'msg': msg})


def login_view(request):
    form = LoginForm(request.POST or None)
    msg = None
    if request.method == 'POST':
        if form.is_valid():
            username = form.cleaned_data.get('username')
            password = form.cleaned_data.get('password')
            user = authenticate(username=username, password=password)
            if user is not None and user.is_admin:
                login(request, user)
                return redirect('adminpage')
            elif user is not None and user.is_docente:
                login(request, user)
                return redirect('docente')
            elif user is not None and user.is_funcionario:
                login(request, user)
                return redirect('funcionario')
            else:
                msg= 'invalid credentials'
        else:
            msg = 'error validating form'
    return render(request, 'login.html', {'form': form, 'msg': msg})


def admin(request):
    return render(request,'admin.html')


def docente(request):
    if request.method == 'POST':
        form = ClaseForm(request.POST)
        if form.is_valid():
            clase_instance = form.save(commit=False)
            clase_instance.usuario = request.user
            unidad_ids = request.POST.getlist('unidades')  # Obtén la lista de IDs de unidades seleccionadas
            contenido_ids = request.POST.getlist('contenidos')  # Obtén la lista de IDs de contenidos seleccionados
            unidades = Unidad.objects.filter(id__in=unidad_ids)  # Filtra las unidades por los IDs seleccionados
            contenidos = Contenido.objects.filter(id__in=contenido_ids)  # Filtra los contenidos por los IDs seleccionados
            clase_instance.save()  # Guarda la instancia de clase primero para obtener el ID
            clase_instance.unidades.set(unidades)  # Asigna las unidades utilizando el método set()
            clase_instance.contenidos.set(contenidos)  # Asigna los contenidos utilizando el método set()
            messages.success(request, '¡Clase registrada exitosamente!')
            return redirect('confirmation_clase')
        else:
            errors = form.errors
            return render(request, 'registrar_clase.html', {'form': form, 'errors': errors})
    else:
        form = ClaseForm()
    return render(request, 'registrar_clase.html', {'form': form})






def view_clases(request):
    clases = Clase.objects.all()
    return render(request, 'clases.html', {'clases': clases})
def funcionario(request):
    clases = Clase.objects.all()
    return render(request, 'clases.html', {'clases': clases})
def confirmation_clase(request):
    # Recuperar la última clase registrada por el usuario actual
    latest_clase = Clase.objects.filter(usuario=request.user).last()

    # Verificar si se encontró una clase
    if latest_clase:
        # Pasar la instancia de la clase como parte del contexto
        return render(request, 'confirmacion_clase.html', {'clase': latest_clase})
    else:
        # Si no se encuentra ninguna clase, manejar el caso en consecuencia
        return render(request, 'confirmacion_clase.html', {'error': 'No se encontraron clases registradas'})

def informes_clases(request):
    # Filtrar las clases por el docente que ha iniciado sesión
    clases_docente = Clase.objects.filter(usuario=request.user)

    # Renderizar la plantilla con las clases del docente
    return render(request, 'informes_clases.html', {'clases': clases_docente})

def obtener_unidades_y_contenidos(request):
    materia_id = request.GET.get('materia_id')
    unidades = Unidad.objects.filter(materia_id=materia_id)
    contenidos = Contenido.objects.filter(unidad__materia_id=materia_id)

    unidades_data = [{'id': unidad.id, 'nombre': unidad.nombre} for unidad in unidades]
    contenidos_data = [{'id': contenido.id, 'nombre': contenido.nombre} for contenido in contenidos]

    data = {'unidades': unidades_data, 'contenidos': contenidos_data}
    return JsonResponse(data)


def buscar_clases(request):
    # Obtener el término de búsqueda del parámetro GET
    query = request.GET.get('q')
    print("Término de búsqueda:", query)

    # Si no hay término de búsqueda, mostrar todas las clases
    if not query:
        clases = Clase.objects.all()
    else:
        # Dividir el término de búsqueda en nombre y apellido
        nombres = query.split()

        # Si se proporcionan tanto el nombre como el apellido en el término de búsqueda
        if len(nombres) == 2:
            # Filtrar las clases por nombre y apellido
            clases = Clase.objects.filter(
                usuario__first_name__icontains=nombres[0],
                usuario__last_name__icontains=nombres[1]
            )
        else:
            # Si no se proporcionan tanto el nombre como el apellido, buscar solo por nombre de usuario
            usuario = User.objects.filter(username=query).first()

            if usuario:
                # Filtrar las clases asociadas al usuario encontrado
                clases = Clase.objects.filter(usuario=usuario)
            else:
                # Si no se encuentra ningún usuario con ese nombre de usuario, devolver un conjunto vacío de clases
                clases = Clase.objects.none()

        print("Clases filtradas:", clases)

    # Renderizar el template con las clases encontradas
    return render(request, 'clases.html', {'clases': clases})
def extract_units_and_contents(text):
    units = []
    contents = {}
    current_unit = None

    # Expresión regular para buscar unidades y contenidos
    unit_pattern = r'^(\d+)\.\s+(.+)$'
    content_pattern = r'^(\d+\.\d+)\.\s+(.+)$'

    # Buscar unidades y contenidos en el texto
    lines = text.split('\n')
    for line in lines:
        # Buscar la unidad
        unit_match = re.match(unit_pattern, line)
        if unit_match:
            unit_number = unit_match.group(1)
            unit_name = unit_match.group(2)
            current_unit = f"{unit_number}. {unit_name}"
            units.append(current_unit)
            contents[current_unit] = []
        # Buscar el contenido
        elif current_unit:
            content_match = re.match(content_pattern, line)
            if content_match:
                content_number = content_match.group(1)
                content_text = content_match.group(2)
                contents[current_unit].append(f"{content_number}. {content_text}")

    return units, contents

def upload_document(request):
    if request.method == 'POST':
        form = DocumentForm(request.POST, request.FILES)
        if form.is_valid():
            # Verificar si se ha adjuntado un archivo
            if 'docfile' in request.FILES:
                doc = request.FILES['docfile']

                # Procesar el archivo Word aquí
                docx_file = docx.Document(doc)
                print(doc)
                # Variables para almacenar el semestre y la materia
                semestre_nombre = None
                materia_nombre = None

                # Buscar el semestre y la materia en el documento
                for para in docx_file.paragraphs:
                    # Buscar patrones en el texto para identificar el semestre y la materia
                    if "Semestre" in para.text:
                        semestre_match = re.search(r'Semestre\s*:\s*([^\n]+)', para.text)
                        if semestre_match:
                            semestre_nombre = semestre_match.group(1).strip()
                            print(f"Semestre encontrado: {semestre_nombre}")
                    elif "Nombre de la Materia" in para.text:
                        materia_match = re.search(r'Nombre de la Materia\s*:\s*([^\n]+)', para.text)
                        if materia_match:
                            materia_nombre = materia_match.group(1).strip()
                            print(f"Materia encontrada: {materia_nombre}")

                # Buscar las unidades y los contenidos en el documento
                text = "\n".join([para.text for para in docx_file.paragraphs])
                print(text)
                units, contents = extract_units_and_contents(text)

                # Imprimir las unidades y sus contenidos
                print("Unidades encontradas:")
                print(units)
                print("Contenidos encontrados:")
                print(contents)

                # Verificar si se encontraron el semestre y la materia
                if semestre_nombre and materia_nombre:
                    # Crear o recuperar instancias de Semestre y Materia y guardarlas en la base de datos
                    semestre, _ = Semestre.objects.get_or_create(nombre=semestre_nombre)
                    materia, _ = Materia.objects.get_or_create(nombre=materia_nombre, semestre=semestre)
                    print("Semestre y materia creados correctamente.")

                    # Guardar las unidades y sus contenidos en la base de datos
                    for unit_text, content_list in contents.items():
                        unit, _ = Unidad.objects.get_or_create(nombre=unit_text, materia=materia)
                        unit.save()  # Guardar la unidad
                        print(f"Unidad creada: {unit}")
                        for content_text in content_list:
                            # Crear una instancia de Contenido para cada contenido encontrado
                            contenido, _ = Contenido.objects.get_or_create(nombre=content_text, unidad=unit)
                            contenido.save()  # Guardar el contenido
                            print(f"Contenido creado: {contenido}")
                    # Redirigir a la página de éxito
                    return HttpResponseRedirect('/success/')
                else:
                    # Manejar el caso en el que no se encontró el semestre o la materia
                    return render(request, 'error.html', {'message': 'No se pudo encontrar el semestre o la materia en el documento.'})
            else:
                # Manejar el caso en el que no se ha adjuntado ningún archivo
                return render(request, 'error.html', {'message': 'No se ha adjuntado ningún archivo.'})
        else:
            # Manejar el caso en el que el formulario no es válido
            return render(request, 'error.html', {'message': 'El formulario no es válido.'})
    else:
        form = DocumentForm()
    return render(request, 'upload_document.html', {'form': form})


from django.http import HttpResponse
from .models import Semestre, Materia, Unidad, Contenido
from docx import Document

from django.http import HttpResponse
from .models import Semestre, Materia, Unidad, Contenido
from docx import Document


from django.contrib import messages


from docx import Document


def ver_asistencia_docente(request):
    hora_inicio = request.GET.get('hora_inicio')
    hora_fin = request.GET.get('hora_fin')

    # Filtrar clases para usuarios con rol de docente
    clases = Clase.objects.filter(usuario__is_docente=True)

    # Aplicar filtros solo si los valores no son None
    if hora_inicio:
        clases = clases.filter(hora_inicio__gte=hora_inicio)
    if hora_fin:
        clases = clases.filter(hora_fin__lte=hora_fin)

    return render(request, 'ver_asistencia.html', {'clases': clases})